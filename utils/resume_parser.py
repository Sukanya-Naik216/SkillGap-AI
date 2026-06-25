from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from PyPDF2 import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_UPLOAD_SIZE_MB = int(os.getenv("MAX_UPLOAD_SIZE_MB", "8"))


class FileValidationError(ValueError):
    """Raised when an uploaded file or text input cannot be processed safely."""


def normalize_whitespace(text: str) -> str:
    """Normalize spacing while preserving line boundaries that help ATS checks."""
    cleaned = text.replace("\x00", " ")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in cleaned.splitlines()]
    return "\n".join(line for line in lines if line)


def validate_filename(filename: str | None) -> str:
    if not filename:
        raise FileValidationError("The uploaded file does not have a valid name.")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise FileValidationError(f"Unsupported file type '{extension}'. Supported types: {supported}.")

    return extension


def validate_file_size(size_bytes: int | None) -> None:
    if size_bytes is None:
        return

    max_size = MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if size_bytes <= 0:
        raise FileValidationError("The uploaded file is empty.")
    if size_bytes > max_size:
        raise FileValidationError(f"File is too large. Maximum size is {MAX_UPLOAD_SIZE_MB} MB.")


def validate_uploaded_file(uploaded_file: Any) -> str:
    """Validate a Streamlit UploadedFile-like object and return its extension."""
    if uploaded_file is None:
        raise FileValidationError("No file was uploaded.")

    extension = validate_filename(getattr(uploaded_file, "name", None))
    file_size = getattr(uploaded_file, "size", None)
    if file_size is None and hasattr(uploaded_file, "getvalue"):
        file_size = len(uploaded_file.getvalue())
    validate_file_size(file_size)
    return extension


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        raise FileValidationError("The PDF could not be opened or is encrypted.") from exc

    text_parts: list[str] = []
    for page in reader.pages:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            # A single problematic page should not discard a whole resume.
            continue

    text = normalize_whitespace("\n".join(text_parts))
    if not text:
        raise FileValidationError("No readable text was found in the PDF. OCR may be required.")
    return text


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise FileValidationError("The DOCX file could not be opened.") from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    # Tables often contain skills and contact details in resumes.
    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)

    text = normalize_whitespace("\n".join(paragraphs))
    if not text:
        raise FileValidationError("No readable text was found in the DOCX file.")
    return text


def extract_text_from_txt_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise FileValidationError("The TXT file encoding could not be detected.")

    text = normalize_whitespace(text)
    if not text:
        raise FileValidationError("No readable text was found in the TXT file.")
    return text


def extract_text_from_upload(uploaded_file: Any) -> str:
    """Extract text from a validated Streamlit upload object."""
    extension = validate_uploaded_file(uploaded_file)
    file_bytes = uploaded_file.getvalue()

    if extension == ".pdf":
        return extract_text_from_pdf_bytes(file_bytes)
    if extension == ".docx":
        return extract_text_from_docx_bytes(file_bytes)
    if extension == ".txt":
        return extract_text_from_txt_bytes(file_bytes)

    raise FileValidationError(f"Unsupported file type: {extension}")


def extract_text_from_path(path: str | Path) -> str:
    """Extract text from a local file path, used by tests and sample utilities."""
    file_path = Path(path)
    extension = validate_filename(file_path.name)
    file_bytes = file_path.read_bytes()
    validate_file_size(len(file_bytes))

    if extension == ".pdf":
        return extract_text_from_pdf_bytes(file_bytes)
    if extension == ".docx":
        return extract_text_from_docx_bytes(file_bytes)
    return extract_text_from_txt_bytes(file_bytes)


def validate_text_input(text: str, label: str, minimum_characters: int = 80) -> None:
    """Reject inputs that are too short to produce a meaningful comparison."""
    if len(normalize_whitespace(text)) < minimum_characters:
        raise FileValidationError(f"{label} text is too short for a reliable analysis.")
